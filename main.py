from loader import extract_from_pdf
from Semantic_splitter import Semantic_split
from Astra_db import vector_database_creation
from prompt_llm import format_docs, generate_response
from collections import defaultdict
from docx import Document
doc = Document()
import pprint

# === Load PDF ===
pdf_path = "Rag_Project\llama.pdf"  # ✅ Update with your actual file path
all_docs = extract_from_pdf(pdf_path)
print("Docs Loadeddd !!!")

pages = Semantic_split(all_docs)
print("\n\nSemantic Search done Successfully")

vector_store = vector_database_creation()
vector_store.add_documents(pages)
print("\n\nPages added Successfully")
# k = int(input("\n\nEnter the k near neighbours : "))
reteriver = vector_store.as_retriever(
    search_type="mmr",
    search_kwargs={'k': 10, 'fetch_k': 30}
)

# relevant_chunks = reteriver.invoke("What is TabelFormer")
# relevant_docs = format_docs(relevant_chunks)

rag_chain = generate_response(retriever=reteriver)
query = input("Enter your query ")
answer=rag_chain.invoke(query)

print("\n\nResponse Generate Suceessfully")

doc.add_heading(f"Q: {query}", level=1)
doc.add_paragraph(answer)
doc.save("rag_output.docx")









